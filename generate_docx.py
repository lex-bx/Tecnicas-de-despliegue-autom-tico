import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
title = doc.add_heading('Informe de Despliegue Automático en Microsoft Azure', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
info = doc.add_table(rows=4, cols=2)
info.cell(0,0).text = 'Asignatura:'
info.cell(0,1).text = 'Técnicas de Despliegue Automático'
info.cell(1,0).text = 'Plataforma:'
info.cell(1,1).text = 'Microsoft Azure (App Service)'
info.cell(2,0).text = 'Aplicación:'
info.cell(2,1).text = 'Task Manager (Python Flask)'
info.cell(3,0).text = 'Repositorio:'
info.cell(3,1).text = 'https://github.com/lex-bx/Tecnicas-de-despliegue-autom-tico'

doc.add_page_break()

# Section 1
doc.add_heading('1. Arquitectura del despliegue', level=1)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Desarrollador')
run.bold = True
pu = p.add_run('\n    | git push\n    v')
doc.add_paragraph('[GitHub Repository]', style='List Bullet')
doc.add_paragraph('    | trigger automático')
doc.add_paragraph('    v')
doc.add_paragraph('[GitHub Actions (CI/CD)]', style='List Bullet')
doc.add_paragraph('    | deploy')
doc.add_paragraph('    v')
doc.add_paragraph('[Azure App Service (Linux)]', style='List Bullet')
doc.add_paragraph('    |-- Gunicorn (WSGI server)')
doc.add_paragraph('    |-- Flask (Aplicación)')
doc.add_paragraph('    |-- SQLite (Base de datos)')

doc.add_paragraph('')
doc.add_heading('Componentes', level=2)

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0,0).text = 'Componente'
table.cell(0,1).text = 'Tecnología'
table.cell(1,0).text = 'Aplicación'
table.cell(1,1).text = 'Python 3.12 + Flask 3.1'
table.cell(2,0).text = 'Servidor WSGI'
table.cell(2,1).text = 'Gunicorn'
table.cell(3,0).text = 'Base de datos'
table.cell(3,1).text = 'SQLite (modo WAL)'
table.cell(4,0).text = 'CI/CD'
table.cell(4,1).text = 'GitHub Actions'
table.cell(5,0).text = 'Hosting'
table.cell(5,1).text = 'Azure App Service (Linux, Plan F1 Gratis)'

# Section 2
doc.add_page_break()
doc.add_heading('2. Creación y configuración de recursos en Azure', level=1)

doc.add_heading('2.1 App Service', level=2)
doc.add_paragraph('Se creó un App Service en Azure Portal con los siguientes parámetros:')

p = doc.add_paragraph()
p.add_run('Nombre: ').bold = True
p.add_run('tecnicas-app-cano')
p = doc.add_paragraph()
p.add_run('Grupo de recursos: ').bold = True
p.add_run('tecnicas-app-cano_group')
p = doc.add_paragraph()
p.add_run('Runtime: ').bold = True
p.add_run('Python 3.12')
p = doc.add_paragraph()
p.add_run('Sistema operativo: ').bold = True
p.add_run('Linux')
p = doc.add_paragraph()
p.add_run('Plan: ').bold = True
p.add_run('F1 (Gratis) - 1 GB de memoria, 60 minutos de CPU/día')
p = doc.add_paragraph()
p.add_run('Región: ').bold = True
p.add_run('East US')
p = doc.add_paragraph()
p.add_run('URL: ').bold = True
p.add_run('https://tecnicas-app-cano-a0fdeqbgfaecfjd4.eastus-01.azurewebsites.net')

doc.add_heading('2.2 Configuración de la aplicación', level=2)
doc.add_paragraph('Se configuró el comando de inicio en Azure Portal:')
p = doc.add_paragraph()
run = p.add_run('gunicorn --bind=0.0.0.0:8000 --timeout 600 --workers 4 app:app')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('2.3 Autenticación básica para despliegue', level=2)
doc.add_paragraph('Se habilitó la autenticación básica de SCM y FTP para permitir el despliegue desde GitHub Actions.')

# Section 3
doc.add_heading('3. Estructura del proyecto', level=1)

table = doc.add_table(rows=10, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0,0).text = 'Archivo'
table.cell(0,1).text = 'Propósito'
table.cell(1,0).text = 'app.py'
table.cell(1,1).text = 'Aplicación Flask'
table.cell(2,0).text = 'requirements.txt'
table.cell(2,1).text = 'Dependencias Python'
table.cell(3,0).text = 'test_app.py'
table.cell(3,1).text = 'Pruebas automáticas'
table.cell(4,0).text = 'templates/index.html'
table.cell(4,1).text = 'Interfaz web'
table.cell(5,0).text = 'static/style.css'
table.cell(5,1).text = 'Estilos'
table.cell(6,0).text = '.github/workflows/deploy.yml'
table.cell(6,1).text = 'CI/CD a Azure App Service'
table.cell(7,0).text = '.github/workflows/deploy-vm.yml'
table.cell(7,1).text = 'CI/CD a Azure VM (alternativa)'
table.cell(8,0).text = 'azure/setup-vm.ps1'
table.cell(8,1).text = 'Script crear VM Azure'
table.cell(9,0).text = 'azure/setup-app.sh'
table.cell(9,1).text = 'Configuración del servidor'

# Section 4
doc.add_page_break()
doc.add_heading('4. Proceso de Integración y Despliegue Continuo (CI/CD)', level=1)

doc.add_heading('4.1 Workflow de GitHub Actions', level=2)
doc.add_paragraph('El archivo .github/workflows/deploy.yml define el pipeline automático que se ejecuta en cada push a la rama main.')

p = doc.add_paragraph()
run = p.add_run('''name: Deploy to Azure App Service

on:
  push:
    branches: [ main, master ]
  workflow_dispatch:

jobs:
  test-and-deploy:
    runs-on: ubuntu-latest
    steps:
      - Checkout del código
      - Configurar Python 3.12
      - Instalar dependencias
      - Ejecutar pruebas
      - Desplegar a Azure App Service''')
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_heading('4.2 Flujo de trabajo', level=2)

flows = [
    'El desarrollador hace "git push" a la rama main',
    'GitHub Actions se activa automáticamente',
    'Se clona el repositorio en un runner de Ubuntu',
    'Se instala Python 3.12 y las dependencias',
    'Se ejecutan las pruebas unitarias (test_app.py)',
    'Si las pruebas pasan, se despliega a Azure App Service',
    'Azure App Service detecta el cambio, reinicia la app y sirve la nueva versión'
]
for f in flows:
    doc.add_paragraph(f, style='List Number')

doc.add_heading('4.3 Secretos de GitHub configurados', level=2)

table = doc.add_table(rows=3, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0,0).text = 'Secreto / Variable'
table.cell(0,1).text = 'Propósito'
table.cell(1,0).text = 'AZURE_PUBLISH_PROFILE'
table.cell(1,1).text = 'Credenciales de despliegue (XML del perfil de publicación)'
table.cell(2,0).text = 'AZURE_APP_NAME'
table.cell(2,1).text = 'Nombre del App Service (tecnicas-app-cano)'

# Section 5
doc.add_heading('5. Funcionamiento de la aplicación', level=1)

doc.add_heading('5.1 Interfaz web', level=2)
doc.add_paragraph('La aplicación permite:')
doc.add_paragraph('Agregar tareas con título y descripción', style='List Bullet')
doc.add_paragraph('Marcar tareas como completadas', style='List Bullet')
doc.add_paragraph('Eliminar tareas', style='List Bullet')

doc.add_heading('5.2 API REST', level=2)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0,0).text = 'Método'
table.cell(0,1).text = 'Ruta'
table.cell(0,2).text = 'Descripción'
table.cell(1,0).text = 'GET'
table.cell(1,1).text = '/api/tasks'
table.cell(1,2).text = 'Listar todas las tareas'
table.cell(2,0).text = 'POST'
table.cell(2,1).text = '/api/tasks'
table.cell(2,2).text = 'Crear tarea'
table.cell(3,0).text = 'PUT'
table.cell(3,1).text = '/api/tasks/{id}'
table.cell(3,2).text = 'Actualizar tarea'
table.cell(4,0).text = 'DELETE'
table.cell(4,1).text = '/api/tasks/{id}'
table.cell(4,2).text = 'Eliminar tarea'

doc.add_heading('5.3 Verificación de funcionamiento', level=2)
doc.add_paragraph('Se verificó que la aplicación responde correctamente desde la URL pública con HTTP 200 y que la API REST opera correctamente.')

# Section 6
doc.add_heading('6. Entregables', level=1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0,0).text = 'Elemento'
table.cell(0,1).text = 'Enlace / Ubicación'
table.cell(1,0).text = 'Repositorio'
table.cell(1,1).text = 'https://github.com/lex-bx/Tecnicas-de-despliegue-autom-tico'
table.cell(2,0).text = 'URL pública'
table.cell(2,1).text = 'https://tecnicas-app-cano-a0fdeqbgfaecfjd4.eastus-01.azurewebsites.net'
table.cell(3,0).text = 'Workflow CI/CD'
table.cell(3,1).text = '.github/workflows/deploy.yml'
table.cell(4,0).text = 'Documentación'
table.cell(4,1).text = 'README.md e informe-despliegue.md'

# Section 7
doc.add_page_break()
doc.add_heading('7. Pasos realizados', level=1)

steps = [
    'Creación de aplicación Flask con interfaz web y API REST (Task Manager)',
    'Inicialización de repositorio Git y subida a GitHub',
    'Creación de App Service en Azure Portal con plan F1 (Gratis)',
    'Habilitación de autenticación básica para publicación',
    'Obtención del perfil de publicación de Azure (.PublishSettings)',
    'Configuración de secretos y variables en GitHub (AZURE_PUBLISH_PROFILE, AZURE_APP_NAME)',
    'Configuración del comando de inicio de gunicorn en Azure',
    'Creación del pipeline CI/CD con GitHub Actions',
    'Push a GitHub activa el despliegue automático',
    'Verificación de la aplicación funcionando en la URL pública'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# Section 8
doc.add_heading('8. Evidencias de funcionamiento', level=1)
doc.add_paragraph('Las siguientes evidencias confirman el correcto funcionamiento del despliegue:')

evidences = [
    'Workflow de GitHub Actions ejecutado exitosamente (estado: Success)',
    'Aplicación responde con HTTP 200 desde la URL pública',
    'API REST operativa: creación y consulta de tareas funcionando',
    'Interfaz web carga correctamente con el título "Task Manager"',
    'Base de datos SQLite funcionando con modo WAL para concurrencia'
]
for e in evidences:
    doc.add_paragraph(e, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('')

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('---')
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run('Elaborado para: Exposición grupal - Técnicas de Despliegue Automático')
run2.italic = True
footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = footer3.add_run('Julio 2026')

doc.save('C:\\Users\\PC\\Desktop\\tecnicas\\informe-despliegue.docx')

print('Word document generated successfully')
